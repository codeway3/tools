#!/usr/bin/env python
# coding: utf-8
# dev-environment: python3.5
# required module: docx, openpyxl
# PS. 部分调试语句没有删除
"""
Usage:
    word_num_counter [-s]

Options:
    -h --help       Show this screen.
    -s              Use cache to run.
"""
import os
import re
import math
import string
import pickle
import datetime
from collections import Counter
import docx
from docx import Document
from docopt import docopt
from openpyxl import Workbook
import jieba
import jieba.analyse as ana
import pandas as pd


SOURCEPATH = './tmp/招股说明书17/'  # 文档目录
MIDDLEPATH = './tmp/storage.pickle'  # 临时存储位置（可忽略
FINPATH = './tmp/风险段落统计.xlsx'  # 统计Excel生成位置
GENPATH = './tmp/风险段落/'  # Word生成位置
STOPWORDS = './tmp/stopword.txt'  # 停用词

dict1 = './tmp/财经金融词汇大全.txt'
dict2 = './tmp/经济 财经 金融 证券 货币 商品 市场 外汇.txt'
jieba.load_userdict(dict1)
jieba.load_userdict(dict2)
pos_dict = './tmp/pos_dic111.txt'
neg_dict = './tmp/neg_dic111.txt'
jieba.load_userdict(pos_dict)
jieba.load_userdict(neg_dict)

# 匹配用的模式串
pattern1_1_in = re.compile(r'^.*重大事项提示.*$')
pattern1_2_in = re.compile(r'^[一二三四五六七八九十]+、.*风险.*')
pattern1_sp = re.compile(r'.*(风险|影响).*')
pattern2_in = re.compile(r'^((第[一二三四五六七八九十]+节)|(第[一二三四五六七八九十]+章))风险因素(及对策)*$')
pattern3_1_in = re.compile(r'^((第[一二三四五六七八九十]+节)|(第[一二三四五六七八九十]+章))管理层讨论与分析$')
pattern3_2_in = re.compile(r'^[一二三四五六七八九十]+、.*(盈利能力.*未来|未来.*盈利能力|盈利能力的|的盈利能力|盈利.*前景|未来趋势).*')

pattern1_1_out = re.compile(r'^(\s*)目(\s*)录(\s*)$')
pattern1_2_out = re.compile(r'^[一二三四五六七八九十]+、.*')
pattern2_out = re.compile(r'^(([一二三四五六七八九十]+、)|(第[一二三四五六七八九十]+节)|(第[一二三四五六七八九十]+章))(发行人|发行人的|公司|本公司)基本情况(\s*)$')
pattern3_1_out = re.compile(r'^((第[一二三四五六七八九十]+节)|(第[一二三四五六七八九十]+章))业务发展.*$')
pattern3_2_out = re.compile(r'^[一二三四五六七八九十]+、.*')


def get_dict_lst_from_file(f):
    d = list()
    if isinstance(f, str):
        f = open(f, 'rb')
        for lineno, ln in enumerate(f, 1):
            line = ln.strip()
            d.append(line.decode('utf-8'))
    return d


# docx 转换为 list
def read_docx(name):
    file = docx.Document(name)
    para_data = list()
    for paragraph in file.paragraphs:
        para_text = paragraph.text
        para_data.append(para_text)
    return para_data


# 文档字数统计，对于字符串列表和单独字符串有不同的处理逻辑
def word_count(d):
    cnt = 0
    if isinstance(d, list):
        for s in d:
            cnt += str_count(s)
    elif isinstance(d, str):
        cnt += str_count(d)
    return cnt


# 字符串字数统计
def str_count(s):
    count_en = count_dg = count_sp = count_zh = count_pu = 0  # 统一将0赋值给这5个变量
    # s_len = len(s)
    for c in s:
        if c in string.ascii_letters:
            count_en += 1
        elif c.isdigit():
            count_dg += 1
        elif c.isspace():
            count_sp += 1
        elif c.isalpha():
            count_zh += 1
        else:
            count_pu += 1
    return count_zh
    # total_chars = count_zh + count_en + count_sp + count_dg + count_pu
    # if total_chars == s_len:
    #     return('总字数：{0},中文字数：{1},英文字数：{2},空格：{3},数字数：{4},标点符号：{5}'.format(s_len, count_zh, count_en, count_sp, count_dg, count_pu))


def getFileName(path):
    filename = []
    f_list = os.listdir(path)
    for i in f_list:
        if os.path.splitext(i)[1] == '.docx':
            filename.append(i)
    return filename


# 暂存字典中的数据（可忽略
def transfer_word(srcdir, tardir):
    ans = list()
    for filename in getFileName(srcdir):
        tmp = dict()
        tmp['name'] = filename
        tmp['text'] = read_docx(srcdir+filename)
        ans.append(tmp)
    file = open(tardir, 'wb')
    pickle.dump(ans, file)
    file.close()
    return ans


# 反序列化暂存数据（可忽略
def deserialize_dict(dir):
    with open(dir, 'rb') as file:
        tmp_dict = pickle.load(file)
    return tmp_dict


# 处理分词词数统计
def words_count(dat):
    df_list = []
    ans = pd.DataFrame([], columns=['word'])
    ana.set_stop_words(STOPWORDS)
    for para in dat:
        wordlist = ana.extract_tags(para, topK=200)
        df = pd.DataFrame(wordlist, columns=['word'])
        df_list.append(df)
    ans = pd.concat(df_list)
    result = ans.shape[0]
    # result = ans.groupby('word').size().sort_values(ascending=False)
    return result
    # 需要控制输出行数时 改这里 改成 return result.head(xx) xx为需要保留的行数


# 情感分数计算 文档i个词 当前词语在该文档中出现j次
def emotion_calc(i, j, wj):
    if j > 0 and wj > 0:
        wij = 1 + math.log(j) * wj
        # print(wij)
    else:
        wij = 0
    return wij


# 数字保留两位小数
def num_fmt(x):
    return int(x*100)/100


# 进行字数统计，并输出到表格中
def generate_xlsx(src_dict):
    pos_dict_lst = get_dict_lst_from_file(pos_dict)
    neg_dict_lst = get_dict_lst_from_file(neg_dict)
    rows = list()
    crows = list()
    emo_rows = list()
    failed_rows = list()
    s_lst = list()
    ss_lst = list()
    word_in_paras_count = [{}, {}, {}]
    word_in_doc_count = {}
    weight_of_word_in_paras = [{}, {}, {}]
    weight_of_word_in_all_doc = {}
    success_doc_num = 0
    for dt in src_dict:
        file_name = dt['name']
        row = list()
        crow = list()
        row.append(file_name)
        crow.append(file_name)
        print(file_name)

        all_num = word_count(dt['text'])
        row.append(all_num)
        crow.append(words_count(dt['text']))

        warning_num = [0, 0, 0]  # 风险字数
        words_num = [0, 0, 0]  # 风险词数
        sumnum = 0
        flag1 = False
        flag3 = False
        flag = [False, False, False]
        paras_to_word = [[], [], []]

        document = Document()
        for para_text in dt['text']:
            para_text = ''.join(para_text.split())
            if re.match(pattern1_1_out, para_text) and flag1:
                flag1 = False
                flag[0] = False
                print(para_text)
            elif re.match(pattern3_1_out, para_text) and flag3:
                flag3 = False
                flag[2] = False
                print(para_text)
            elif re.match(pattern1_2_out, para_text) and flag1 and flag[0]:
                flag[0] = False
                print(para_text)
            elif re.match(pattern2_out, para_text) and flag[1]:
                flag[1] = False
                print(para_text)
            elif re.match(pattern3_2_out, para_text) and flag3 and flag[2]:
                flag[2] = False
                print(para_text)

            if flag[0]:
                warning_num[0] += word_count(para_text)
                words_num[0] += words_count([para_text])
                paras_to_word[0].append(para_text)
            elif flag[1]:
                warning_num[1] += word_count(para_text)
                words_num[1] += words_count([para_text])
                paras_to_word[1].append(para_text)
            elif flag[2]:
                warning_num[2] += word_count(para_text)
                words_num[2] += words_count([para_text])
                paras_to_word[2].append(para_text)
            elif flag1 and re.match(pattern1_sp, para_text):
                warning_num[0] += word_count(para_text)
                words_num[0] += words_count([para_text])
                paras_to_word[0].append(para_text)

            if flag[0] or flag[1] or flag[2] or (flag1 and re.match(pattern1_sp, para_text)):
                sumnum += word_count(para_text)

            if re.match(pattern1_2_in, para_text) and flag1:
                flag[0] = True
                print(para_text)
            elif re.match(pattern2_in, para_text):
                flag1 = False
                flag[1] = True
                print(para_text)
            elif re.match(pattern3_2_in, para_text) and flag3:
                flag[2] = True
                print(para_text)
            elif re.match(pattern1_1_in, para_text):
                flag1 = True
                print(para_text)
            elif re.match(pattern3_1_in, para_text):
                flag3 = True
                print(para_text)
        row.extend(warning_num)
        row.append(sum(warning_num))
        words_sum = sum(words_num)
        crow.extend(words_num)
        crow.append(words_sum)
        if sumnum/all_num > 0.8 or warning_num[1] == 0 or warning_num[2] == 0:
            failed_rows.append(row)
        else:
            success_doc_num += 1
            rows.append(row)
            crows.append(crow)
            # 统计情感分数
            ss = ''
            s = ['', '', '']
            for i in range(3):
                s[i] = ' '.join(paras_to_word[i])
                c = Counter(jieba.cut(s[i]))
                for wd in c:
                    if c[wd] > 0:
                        cnt = word_in_paras_count[i].setdefault(wd, 0)
                        word_in_paras_count[i][wd] = cnt + 1
            ss = s[0] + s[1] + s[2]
            s_lst.append(s)
            ss_lst.append(ss)
            cc = Counter(jieba.cut(ss))
            for wd in cc:
                if cc[wd] > 0:
                    cnt = word_in_doc_count.setdefault(wd, 0)
                    word_in_doc_count[wd] = cnt + 1
            # 生成.docx
            for paras in paras_to_word:
                for para in paras:
                    if para:
                        document.add_paragraph(para)
            document.save(GENPATH+file_name)
        print()

    for i in range(3):
        for wd in word_in_paras_count[i]:
            weight_of_word_in_paras[i][wd] = math.log(success_doc_num/word_in_paras_count[i][wd])
    for wd in word_in_doc_count:
        weight_of_word_in_all_doc[wd] = math.log(success_doc_num/word_in_doc_count[wd])
    # print(word_in_doc_count)
    # print(weight_of_word_in_all_doc)
    for num, ss in enumerate(ss_lst):
        cc = Counter(jieba.cut(ss))
        weight_pos = 0
        weight_neg = 0
        for wd in cc:
            if wd in pos_dict_lst:
                weight_pos += emotion_calc(crows[num][5], cc[wd], weight_of_word_in_all_doc.setdefault(wd, 0))
            if wd in neg_dict_lst:
                weight_neg += emotion_calc(crows[num][5], cc[wd], weight_of_word_in_all_doc.setdefault(wd, 0))
        weight_pos = 1/(1+math.log(crows[num][5]))*weight_pos
        weight_neg = 1/(1+math.log(crows[num][5]))*weight_neg
        emo_row = [crows[num][0], num_fmt(weight_pos), num_fmt(weight_neg)]
        for i in range(3):
            weight_pos = 0
            weight_neg = 0
            cc = Counter(jieba.cut(s_lst[num][i]))
            for wd in cc:
                if wd in pos_dict_lst:
                    weight_pos += emotion_calc(crows[num][i+2], cc[wd], weight_of_word_in_paras[i].setdefault(wd, 0))
                if wd in neg_dict_lst:
                    weight_neg += emotion_calc(crows[num][i+2], cc[wd], weight_of_word_in_paras[i].setdefault(wd, 0))
            emo_row.append(num_fmt(weight_pos))
            emo_row.append(num_fmt(weight_neg))
        emo_rows.append(emo_row)
    print(emo_rows)

    # 创建Excel文档
    wb = Workbook()
    # 操作工作表1 风险段落字数统计
    sheet1 = wb.active
    sheet1.title = '主要风险段落字数统计'
    sheet1.append(['文档名字', '总字数', '重大事项提示-风险', '风险因素', '管理层讨论与分析-未来盈利能力', '风险段落总字数'])
    for row in rows:
        sheet1.append(row)
    # 操作工作表2 风险段落词数统计
    wb.create_sheet('主要风险段落词数统计', index=1)
    sheet2 = wb['主要风险段落词数统计']
    sheet2.append(['文档名字', '总词数', '重大事项提示-风险', '风险因素', '管理层讨论与分析-未来盈利能力', '风险段落总词数'])
    for crow in crows:
        sheet2.append(crow)
    # 操作工作表3 情感打分
    wb.create_sheet('文档情感打分', index=2)
    sheet3 = wb['文档情感打分']
    sheet3.append(['文档名字',
                   '总体积极分数', '总体消极分数',
                   '重大事项提示-风险积极分数', '重大事项提示-风险消极分数',
                   '风险因素积极分数', '风险因素消极分数',
                   '管理层讨论与分析-未来盈利能力积极分数', '管理层讨论与分析-未来盈利能力消极分数'])
    for emo_row in emo_rows:
        sheet3.append(emo_row)
    # 操作工作表4 处理失败文档 数字为对应的字数统计结果
    wb.create_sheet('处理失败文档', index=3)
    sheet4 = wb['处理失败文档']
    for row in failed_rows:
        sheet4.append(row)
    # 保存Excel文档
    wb.save(FINPATH)


if __name__ == '__main__':
    starttime = datetime.datetime.now()
    args = docopt(__doc__)
    if args['-s']:
        fd = deserialize_dict(MIDDLEPATH)
    else:
        fd = transfer_word(SOURCEPATH, MIDDLEPATH)
    generate_xlsx(fd)
    endtime = datetime.datetime.now()
    print('处理完成，总共用时 {}'.format(endtime - starttime))
